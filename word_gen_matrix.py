"""
word_gen_matrix.py — Integrations matrix table, generator-served system expansion,
and diagram image insertion.
Change this file for bugs in the matrix table layout or generator-served section.
"""

from copy import deepcopy
from lxml import etree
from docx.shared import Inches

from defaults import MATRIX_DEFAULTS
from constants import SYSTEMS
from word_gen_replacements import build_replacements


def populate_matrix_table(doc, data):
    """
    Find the integrations matrix table and expand each system row into
    one row per integration sub-type, with System A/B cells vMerge-spanned.
    pre_action is handled as TWO dynamic sections (pre_pan + pre_sys), identical
    in mechanism to sprinkler, standpipe, etc.
    """

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = f"{{{WNS}}}"

    systems_data = data.get("systems", {})

    # Compute gen_class display label for matrix System B cell
    _gen_class_raw = systems_data.get("generator", {}).get("gen_class", "non-emergency")
    _gen_class_display = "Emergency" if _gen_class_raw == "emergency" else "Non-Emergency"

    SYS_B_LABELS = {
        "sprinkler": "Sprinkler System",
        "standpipe": "Standpipe System",
        "fire_pump": "Fire Pump",
        "generator": f"{_gen_class_display} Generator",
        "maglock": "Maglocks",
        "door_holders": "Door Holders",
        "ahu": "AHU/Fan Shutdown",
        "smoke_dampers": "Smoke Dampers",
        "fire_shutters": "Fire Shutters",
        "kitchen_hood": "Kitchen Hood Suppression System",
        "water_mist": "Water Mist System",
        "elevator": "Elevator",
    }

    # Build sections: (sys_b_label, rows, sys_a_override)
    # sys_a_override is None for all systems where System A = "Fire Alarm"
    # For pre_sys, System A = "Pre-Action Sprinkler System Panel"
    sections = []

    # Monitoring — always present, under fire_alarm
    fa_sys = systems_data.get("fire_alarm", {})
    mon_rows = fa_sys.get("matrix_mon", [])
    if not mon_rows:
        mon_rows = [{"integration": r[0], "normal_mode": r[1], "fire_mode": r[2]}
                    for r in MATRIX_DEFAULTS.get("fire_alarm_monitoring", [])]
    sections.append(("Monitoring Station", mon_rows, None))

    for sys_info in SYSTEMS:
        key = sys_info["key"]
        if key == "fire_alarm":
            continue
        sys_d = systems_data.get(key, {})
        if not sys_d.get("present", False):
            continue

        if key == "pre_action":
            _has_panel = sys_d.get("pre_action_panel", False)
            # preac_pan_or_fa: matches replace_all value used in body/headings
            _preac = "Pre-Action Panel" if _has_panel else "Fire Alarm"

            # pre_pan: Fire Alarm -> Pre-Action Sprinkler System Panel
            # Only exists when a designated pre-action panel is present
            if _has_panel:
                pap_rows = sys_d.get("pap_matrix_rows", [])
                if not pap_rows:
                    pap_rows = [{"integration": r[0], "normal_mode": r[1], "fire_mode": r[2]}
                                for r in MATRIX_DEFAULTS.get("pre_action_panel", [])]
                sections.append(("Pre-Action Sprinkler System Panel", pap_rows, None))
                print(f"[MATRIX] pre_pan: {len(pap_rows)} rows")

            # pre_sys: preac_pan_or_fa -> Pre-Action Sprinkler System
            # System A = "Pre-Action Panel" (panel present) or "Fire Alarm" (no panel)
            pa_rows = sys_d.get("matrix_rows", [])
            if not pa_rows:
                pa_rows = [{"integration": r[0], "normal_mode": r[1], "fire_mode": r[2]}
                           for r in MATRIX_DEFAULTS.get("pre_action", [])]
            # sys_a_override=None when no panel (System A stays "Fire Alarm" from template_tr)
            sections.append(("Pre-Action Sprinkler System", pa_rows,
                             _preac if _has_panel else None))
            print(f"[MATRIX] pre_sys: {len(pa_rows)} rows, sys_a={_preac!r}")
            continue

        rows = sys_d.get("matrix_rows", [])
        if not rows:
            rows = [{"integration": r[0], "normal_mode": r[1], "fire_mode": r[2]}
                    for r in MATRIX_DEFAULTS.get(key, [])]
        sections.append((SYS_B_LABELS.get(key, key), rows, None))
        print(f"[MATRIX] {key}: {len(rows)} rows")

    # Find matrix table
    matrix_table = None
    for table in doc.tables:
        for row in table.rows:
            if any("System A" in cell.text for cell in row.cells):
                matrix_table = table
                break
        if matrix_table:
            break
    if matrix_table is None:
        print("[MATRIX] ERROR: matrix table not found!")
        return
    print(f"[MATRIX] found matrix table with {len(matrix_table.rows)} rows")

    tbl_el = matrix_table._tbl

    # Save header row and gen_con template row; remove all data rows
    data_rows = list(matrix_table.rows[1:])
    template_tr = deepcopy(data_rows[0]._tr) if data_rows else None
    # gen_con row: first Emergency Power row in template (data_rows[7] = row 8 overall)
    template_tr_gen_con = deepcopy(data_rows[7]._tr) if len(data_rows) > 7 else None
    for row in data_rows:
        tbl_el.remove(row._tr)
    if template_tr is None:
        print("[MATRIX] ERROR: no template row to clone from")
        return

    def _clean_cell_text(tc, text):
        text = text.replace("\n", " ").strip()
        paragraphs = tc.findall(f"{W}p")
        if paragraphs:
            p = paragraphs[0]
            for r in p.findall(f".//{W}r"):
                for t in r.findall(f"{W}t"):
                    t.text = ""
            runs = p.findall(f".//{W}r")
            if runs:
                t_els = runs[0].findall(f"{W}t")
                if t_els:
                    t_els[0].text = text
                    if text:
                        t_els[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            for p_extra in paragraphs[1:]:
                tc.remove(p_extra)
            for r in p.findall(f"{W}r"):
                t_els = r.findall(f"{W}t")
                if not t_els or all((t.text or "").strip() == "" for t in t_els):
                    if r != (p.findall(f"{W}r") or [None])[0]:
                        p.remove(r)

    def _set_vmerge(tc, restart=False):
        tcp = tc.find(f"{W}tcPr")
        if tcp is None:
            tcp = etree.SubElement(tc, f"{W}tcPr")
            tc.insert(0, tcp)
        for vm in tcp.findall(f"{W}vMerge"):
            tcp.remove(vm)
        vm = etree.SubElement(tcp, f"{W}vMerge")
        if restart:
            vm.set(f"{W}val", "restart")

    def _ensure_empty_para(tc):
        paragraphs = tc.findall(f"{W}p")
        if paragraphs:
            p = paragraphs[0]
            for r in p.findall(f".//{W}r"):
                for t in r.findall(f"{W}t"):
                    t.text = ""
            for p_extra in paragraphs[1:]:
                tc.remove(p_extra)

    for sys_b_label, int_rows, sys_a_override in sections:
        n = len(int_rows)
        for i, row_data in enumerate(int_rows):
            new_tr = deepcopy(template_tr)
            tcs = new_tr.findall(f".//{W}tc")

            if i == 0:
                # Override System A if needed (e.g. pre_sys rows)
                if sys_a_override:
                    _clean_cell_text(tcs[0], sys_a_override)
                _clean_cell_text(tcs[1], sys_b_label)
                if n > 1:
                    _set_vmerge(tcs[0], restart=True)
                    _set_vmerge(tcs[1], restart=True)
            else:
                _ensure_empty_para(tcs[0])
                _ensure_empty_para(tcs[1])
                _set_vmerge(tcs[0], restart=False)
                _set_vmerge(tcs[1], restart=False)

            _clean_cell_text(tcs[2], row_data.get("integration", "").strip())
            _clean_cell_text(tcs[3], row_data.get("normal_mode", "").strip())
            _clean_cell_text(tcs[4], row_data.get("fire_mode", "").strip())

            tbl_el.append(new_tr)

    # Generator connected systems rows (emergency generator only)
    if template_tr_gen_con is not None:
        gen_sys = systems_data.get("generator", {})
        is_emergency = gen_sys.get("gen_class", "non-emergency") == "emergency"
        gen_present = gen_sys.get("present", False)

        GEN_CON_MAP = {
            "Fire Alarm": ("Fire Alarm System", "gen_con_fa"),
            "Fire Pump": ("Fire Pump", "gen_con_fpmp"),
            "Maglocks": ("Maglocks", "gen_con_mglck"),
            "Door Holders": ("Door Holders", "gen_con_dhldr"),
            "AHU/Fan": ("Air Handling Units", "gen_con_ahu"),
            "Smoke Dampers": ("Smoke Dampers", "gen_con_sdmpr"),
            "Fire Shutters": ("Fire Shutters", "gen_con_fshtr"),
            "Kitchen Hood": ("Kitchen Hood Suppression System", "gen_con_ktchn"),
            "Water Mist": ("Water Mist System", "gen_con_watmst"),
            "Elevator": ("Elevators", "gen_con_elev"),
        }

        replacements = build_replacements(data)

        if is_emergency and gen_present:
            served_labels = gen_sys.get("gen_served", [])
            served_rows = gen_sys.get("gen_served_matrix_rows", [])
            served_lookup = {}
            for r in served_rows:
                integ = r.get("integration", "").strip()
                served_lookup[integ] = r

            n_served = len(served_labels)
            for i, lbl in enumerate(served_labels):
                if lbl not in GEN_CON_MAP:
                    # Custom user-added system — use generic gen-served text
                    from defaults import GEN_SERVED_NORMAL, GEN_SERVED_GENMODE
                    sys_b_name = lbl
                    normal_val = GEN_SERVED_NORMAL.replace("{{connected_system}}", lbl)
                    fire_val   = GEN_SERVED_GENMODE.replace("{{connected_system}}", lbl)
                else:
                    sys_b_name, prefix = GEN_CON_MAP[lbl]
                    normal_val = replacements.get(f"{prefix}_normal_mode", "")
                    fire_val   = replacements.get(f"{prefix}_fire_mode", "")

                new_tr = deepcopy(template_tr_gen_con)
                tcs = new_tr.findall(f".//{W}tc")

                if i == 0:
                    full_a = "".join(t.text or "" for t in tcs[0].findall(f".//{W}t"))
                    repl_a = full_a
                    for ph, val in replacements.items():
                        repl_a = repl_a.replace(f"{{{{{ph}}}}}", str(val) if val else "")
                    _clean_cell_text(tcs[0], repl_a)
                    if n_served > 1:
                        _set_vmerge(tcs[0], restart=True)
                else:
                    _ensure_empty_para(tcs[0])
                    _set_vmerge(tcs[0], restart=False)

                _clean_cell_text(tcs[1], sys_b_name)
                _clean_cell_text(tcs[3], normal_val)
                _clean_cell_text(tcs[4], fire_val)
                tbl_el.append(new_tr)


def expand_gen_served_system(doc, data):
    """
    Find the {{gen_served_system}} placeholder paragraph in any table cell,
    clone it once per served system with bullet formatting.
    Also expand {{gen_served_list}} in body paragraphs into separate bulleted paragraphs.
    """
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = f"{{{WNS}}}"

    gen_sys = data.get("systems", {}).get("generator", {})
    served_labels = gen_sys.get("gen_served", [])
    GEN_DISPLAY = {
        "Fire Alarm": "Fire Alarm System",
        "Fire Pump": "Fire Pump", "Maglocks": "Electromagnetic Locks",
        "Door Holders": "Door Holders", "AHU/Fan": "Air Handling Units",
        "Smoke Dampers": "Smoke Dampers", "Fire Shutters": "Fire Shutters",
        "Kitchen Hood": "Kitchen Hood Suppression System",
        "Water Mist": "Water Mist System", "Elevator": "Elevator",
    }
    served_names = [GEN_DISPLAY.get(lbl, lbl) for lbl in served_labels]
    if not served_names:
        served_names = [""]

    placeholder = "{{gen_served_system}}"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para_el in list(cell._element.findall(f".//{W}p")):
                    full = "".join(t.text or "" for t in para_el.findall(f".//{W}t"))
                    if placeholder not in full:
                        continue
                    parent = para_el.getparent()
                    insert_idx = list(parent).index(para_el)
                    # Insert one paragraph per served system, each on its own line
                    for i, name in enumerate(served_names):
                        new_p = deepcopy(para_el)
                        # Remove all existing runs
                        for r in new_p.findall(f"{W}r"):
                            new_p.remove(r)
                        # Ensure paragraph properties with indent
                        pPr = new_p.find(f"{W}pPr")
                        if pPr is None:
                            pPr = etree.SubElement(new_p, f"{W}pPr")
                            new_p.insert(0, pPr)
                        ind = pPr.find(f"{W}ind")
                        if ind is None:
                            ind = etree.SubElement(pPr, f"{W}ind")
                        ind.set(f"{W}left", "360")
                        # Add a single run with the bullet + name
                        r_el = etree.SubElement(new_p, f"{W}r")
                        # Copy rPr from original if present
                        orig_runs = para_el.findall(f"{W}r")
                        if orig_runs:
                            orig_rPr = orig_runs[0].find(f"{W}rPr")
                            if orig_rPr is not None:
                                r_el.insert(0, deepcopy(orig_rPr))
                        t_el = etree.SubElement(r_el, f"{W}t")
                        t_el.text = f"•  {name}"
                        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                        parent.insert(insert_idx + i, new_p)
                    parent.remove(para_el)
                    break

    list_ph = "{{gen_served_list}}"
    for para in doc.paragraphs:
        if list_ph in para.text:
            placeholder_el = para._element
            parent = placeholder_el.getparent()
            insert_idx = list(parent).index(placeholder_el)
            for i, name in enumerate(served_names):
                new_el = deepcopy(placeholder_el)
                # Remove all runs and replace with a single bullet run
                for r in new_el.findall(f"{W}r"):
                    new_el.remove(r)
                pPr = new_el.find(f"{W}pPr")
                if pPr is None:
                    pPr = etree.SubElement(new_el, f"{W}pPr")
                    new_el.insert(0, pPr)
                ind = pPr.find(f"{W}ind")
                if ind is None:
                    ind = etree.SubElement(pPr, f"{W}ind")
                ind.set(f"{W}left", "360")
                r_el = etree.SubElement(new_el, f"{W}r")
                orig_runs = placeholder_el.findall(f"{W}r")
                if orig_runs:
                    orig_rPr = orig_runs[0].find(f"{W}rPr")
                    if orig_rPr is not None:
                        r_el.insert(0, deepcopy(orig_rPr))
                t_el = etree.SubElement(r_el, f"{W}t")
                t_el.text = f"•  {name}"
                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                parent.insert(insert_idx + i, new_el)
            parent.remove(placeholder_el)
            break


def insert_diagram_image(doc, png_path):
    """Replace {{interconnection_diag_png}} placeholder paragraph with the diagram image."""
    for para in doc.paragraphs:
        if "{{interconnection_diag_png}}" in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(png_path, width=Inches(6.0))
            return